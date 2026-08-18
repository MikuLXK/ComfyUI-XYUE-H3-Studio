"""Safe, bounded reference-document parsing for prompt enhancement."""

from __future__ import annotations

import json
import os
from pathlib import Path
from typing import Any

try:
    from ..core.contracts import DOCUMENT_EXTENSIONS
except ImportError:  # direct test import from the plugin root
    from core.contracts import DOCUMENT_EXTENSIONS

MAX_FILE_BYTES = 25 * 1024 * 1024


def validate_document_path(path: str | os.PathLike[str], root: str | os.PathLike[str] | None = None) -> Path:
    candidate = Path(path).resolve()
    if root is not None:
        boundary = Path(root).resolve()
        if candidate != boundary and boundary not in candidate.parents:
            raise ValueError("文档路径超出允许目录")
    if candidate.suffix.lower() not in DOCUMENT_EXTENSIONS:
        raise ValueError(f"不支持的文档类型：{candidate.suffix}")
    if not candidate.is_file():
        raise ValueError(f"文档不存在：{candidate.name}")
    if candidate.stat().st_size > MAX_FILE_BYTES:
        raise ValueError("文档超过 25MB 限制")
    return candidate


def extract_text(path: str | os.PathLike[str], limit: int = 80_000) -> tuple[str, dict[str, Any]]:
    file_path = validate_document_path(path)
    suffix = file_path.suffix.lower()
    if suffix in {".txt", ".md"}:
        text = file_path.read_text(encoding="utf-8", errors="replace")
    elif suffix == ".json":
        text = json.dumps(json.loads(file_path.read_text(encoding="utf-8", errors="replace")), ensure_ascii=False, indent=2)
    elif suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("解析 PDF 需要安装 pypdf") from exc
        reader = PdfReader(str(file_path))
        text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
    elif suffix == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("解析 DOCX 需要安装 python-docx") from exc
        document = Document(str(file_path))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        for table in document.tables:
            text += "\n" + "\n".join(" | ".join(cell.text for cell in row.cells) for row in table.rows)
    else:
        raise ValueError(f"不支持的文档类型：{suffix}")
    text = text.strip()
    clipped = len(text) > limit
    return text[:limit], {"filename": file_path.name, "characters": len(text), "clipped": clipped}
